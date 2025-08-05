from docx import Document
import io
import re


# 📝 Converts a structured summary JSON into a formatted in-memory DOCX file
def create_docx_in_memory(summary_json, document_title):
    
    # Create a new Word document
    doc = Document()
    doc.add_heading(document_title, level=0)     # Add title as top-level heading

    # 📄 Loop through all sections in the summary JSON
    for section in summary_json.get("sections", []):
        doc.add_heading(section["heading"], level=1)    # Add section heading
        
        # 📝 Split section content by line (assuming newline-separated bullet points)
        for line in section["content"].split("\n"):
            line = line.strip()
            
            # ➤ Format as a bullet point if it starts with '- '
            if line.startswith("- "):
                line = line[2:].strip()
                para = doc.add_paragraph(style="List Bullet")
                parts = re.split(r"(\*\*.*?\*\*)", line)
                
                # 🎯 Split and apply bold formatting to any parts enclosed in **
                for part in parts:
                    run = para.add_run()
                    if part.startswith("**") and part.endswith("**"):
                        run.text = part[2:-2]
                        run.bold = True
                    else:
                        run.text = part
            
            # ➕ Add any regular lines not formatted as bullets
            else:
                doc.add_paragraph(line.strip())

    # 💾 Save document to an in-memory BytesIO stream
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # 📤 Return the binary stream (for upload or download)
    return doc_stream