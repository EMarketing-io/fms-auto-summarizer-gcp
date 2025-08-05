from docx import Document
import io


# 📝 Generates a structured DOCX meeting summary from the provided summary data
def generate_docx(summary_data, company_name, meeting_date):
    
    # Create a new Word document
    doc = Document()

    # 📌 Add main title and meeting date
    doc.add_heading(f"{company_name} Meeting Notes", level=0)
    doc.add_paragraph(f"Date: {meeting_date}", style="Heading 2").alignment = 2
    
    # 🗒️ Section 1: Minutes of the Meeting (MoM)
    doc.add_heading("1. Minutes of the Meeting (MoM)", level=1)
    for line in summary_data["mom"]:
        doc.add_paragraph(line.strip(), style="List Bullet")
    
    # ✅ Section 2: To-Do List
    doc.add_heading("2. To-Do List", level=1)
    for item in summary_data["todo_list"]:
        doc.add_paragraph(item.strip(), style="List Bullet")

    # 📌 Section 3: Action Plan with subcategories
    doc.add_heading("3. Action Points / Action Plan", level=1)

    # Mapping of internal keys to human-readable section titles
    section_titles = {
        "decision_made": "Key Decisions Made",
        "key_services_to_promote": "Key Services to Promote",
        "target_geography": "Target Geography",
        "budget_and_timeline": "Budget and Timeline",
        "lead_management_strategy": "Lead Management Strategy",
        "next_steps_and_ownership": "Next Steps and Ownership",
    }

    # ➕ Populate each sub-section under Action Plan
    for key, title in section_titles.items():
        doc.add_heading(title, level=2)
        for item in summary_data["action_plan"].get(key, []):
            doc.add_paragraph(item.strip(), style="List Bullet")

    # 📤 Convert the completed DOCX document into a binary stream
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)

    # Return the document content as binary data (for in-memory uploads)
    return docx_stream.read()